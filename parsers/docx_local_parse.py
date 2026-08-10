"""DOCX 로컬 결정론 파서 (docparse Tier 0): python-docx로 외부 API 없이 추출.

다른 *_parse.py와 동일한 인터페이스:
  python docx_local_parse.py "<파일.docx>"        → <파일>_docxlocal.md

DOCX는 단락·표 셀·헤딩 스타일이 XML(document.xml)에 명시된 포맷이라 파싱이
추론이 아니라 읽기다(hwpx-tomd와 같은 철학). 본문 블록(헤딩·단락·표)을 문서
순서대로 보존하고, 병합 셀은 앵커에 1회만 기록 + 병합 고지. 머리글·바닥글은
본문과 별개 XML 파트이므로 파트별로 같은 추출(표·첫/짝수 페이지 변형 포함)과
recall 검증을 적용해 본문 뒤에 라벨을 붙여 표기한다.

전수 recall 검증 내장: 각 XML 파트의 모든 텍스트(w:t, 탭·줄바꿈은 공백으로
합류)를 가장 가까운 단락(w:p) 단위로 자체 집계(zipfile + ElementTree,
python-docx와 독립)해 추출 결과와 토큰 멀티셋 양방향 대조한다. python-docx가
구조적으로 못 읽는 텍스트박스(w:pict/w:txbxContent)·필드 텍스트가 있으면
반드시 누락으로 드러나며, 이때 출력 파일을 만들지 않는다 — 기존 티어(Upstage·
LlamaParse)로 승격할 것. 각주·미주에 실제 텍스트가 있으면 같은 이유로
거부한다(python-docx 미지원 영역).

한계: 이미지 안의 텍스트는 범위 밖(존재 시 경고, 출력은 유지 — hwpx_local과
동일 정책). 수식(OMML, m:t)은 양쪽 관점 모두 다루지 않아 검증으로도 잡히지
않는다(수식 문서는 기존 티어 권장). 심볼 글리프(w:sym)도 같은 대칭 유실
계급이라 존재 감지 시 거부한다. 중첩 표는 마크다운으로 구조 보존이 불가해
거부한다. DOC(구형 바이너리)는 Word/LibreOffice로 DOCX 저장 후 입력.
"""
import os
import re
import sys
import traceback
import zipfile
import xml.etree.ElementTree as ET
from collections import Counter

# Windows cp949 안전 (다른 파서와 동일 패턴: 양쪽 스트림 + replace)
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

SUPPORTED_EXTENSIONS = [".docx"]
MAX_PROBLEM_SAMPLES = 10
NS_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
HEADING_RE = re.compile(r"^(?:Heading|제목)\s*(\d)$")


def get_output_filename(input_file):
    """입력 파일 경로 기반 출력 경로(<base>_docxlocal.md) 생성."""
    dir_path = os.path.dirname(input_file)
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    output_name = f"{base_name}_docxlocal.md"
    return os.path.join(dir_path, output_name) if dir_path else output_name


def _cell_md(text):
    return " ".join(text.split()).replace("|", "\\|")


def _grid_to_markdown(grid):
    if not grid:
        return ""
    width = max(len(row) for row in grid)
    lines = []
    for i, row in enumerate(grid):
        cells = [_cell_md(v) for v in row] + [""] * (width - len(row))
        lines.append("| " + " | ".join(cells) + " |")
        if i == 0:
            lines.append("|" + " --- |" * width)
    return "\n".join(lines)


def _xml_paragraph_texts(xml_bytes):
    """XML의 모든 w:t를 '가장 가까운 조상 w:p' 단위로 연결한 문자열 목록.

    run 분절(한 단어가 여러 w:t로 쪼개져 저장됨)을 단락 단위 연결로 흡수하고,
    텍스트박스 안의 중첩 w:p는 자기 단락으로 귀속되므로 본문과 섞이지 않는다.
    python-docx를 쓰지 않는 독립 관점이다.
    """
    root = ET.fromstring(xml_bytes)
    parent = {c: p for p in root.iter() for c in p}
    texts = {}
    order = []
    # python-docx의 Run.text 번역 규칙(t·tab·br·cr·noBreakHyphen·ptab)과 토큰
    # 단위로 대칭이 되게 합류시킨다. 구분자 없이 w:t만 이어붙이면 소프트
    # 줄바꿈·탭이 있는 흔한 문서가 "첫줄둘째줄"처럼 뭉쳐 전부 오거부된다.
    text_like = {
        f"{NS_W}t", f"{NS_W}tab", f"{NS_W}br", f"{NS_W}cr",
        f"{NS_W}noBreakHyphen", f"{NS_W}ptab",
    }
    for el in root.iter():
        if el.tag not in text_like:
            continue
        node = el
        anchor = root
        while node in parent:
            node = parent[node]
            if node.tag == f"{NS_W}p":
                anchor = node
                break
        if anchor not in texts:
            texts[anchor] = []
            order.append(anchor)
        if el.tag == f"{NS_W}t":
            piece = el.text or ""
        elif el.tag == f"{NS_W}noBreakHyphen":
            piece = "-"  # python-docx는 하이픈 문자로 번역한다
        elif el.tag == f"{NS_W}br" and el.get(f"{NS_W}type") in ("page", "column"):
            piece = ""  # python-docx는 페이지·단 나눔을 빈 문자열로 번역한다
        else:
            piece = " "
        texts[anchor].append(piece)
    return ["".join(texts[p]) for p in order]


def _tokens(strings):
    return Counter(tok for s in strings for tok in s.split())


def _part_has_text(zf, name):
    """해당 XML 파트에 실제 w:t 텍스트가 있는지 (각주·미주 감지용)."""
    try:
        data = zf.read(name)
    except KeyError:
        return False
    return any(s.strip() for s in _xml_paragraph_texts(data))


def _heading_level(paragraph):
    style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
    if style_name in ("Title", "제목"):
        return 1
    m = HEADING_RE.match(style_name)
    if m:
        return min(int(m.group(1)), 6)
    return None


def extract_blocks(doc):
    """본문을 문서 순서대로 (마크다운 블록 목록, 원문 문자열 목록, 병합 수, 중첩표 수)로.

    원문 문자열 목록은 recall 검증용(마크다운 이스케이프 이전의 원 텍스트).
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    blocks = []
    raw_texts = []
    merged_cells = 0
    nested_tables = 0

    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            raw_texts.append(item.text)
            level = _heading_level(item)
            if level:
                blocks.append("#" * level + " " + " ".join(text.split()))
            else:
                blocks.append(text)
        elif isinstance(item, Table):
            grid = []
            # 값이 아니라 참조를 보관한다: lxml 프록시는 참조가 사라지면 GC 후
            # id가 재사용될 수 있어, id만 저장하면 정상 셀을 병합으로 오인한다.
            seen_tc = {}
            for row in item.rows:
                row_vals = []
                for cell in row.cells:
                    if cell.tables:
                        nested_tables += 1
                    tc = cell._tc
                    key = id(tc)
                    if key in seen_tc:
                        merged_cells += 1
                        row_vals.append("")
                        continue
                    seen_tc[key] = tc
                    text = " ".join(p.text for p in cell.paragraphs)
                    raw_texts.append(text)
                    row_vals.append(text)
                grid.append(row_vals)
            md = _grid_to_markdown(grid)
            if md:
                blocks.append(md)
    return blocks, raw_texts, merged_cells, nested_tables


def _verify_tokens(xml_bytes, raw_texts, label):
    """추출 원문 vs XML 자체 집계를 토큰 멀티셋 양방향 대조.

    반환: (XML 토큰 수, 문제 목록[str])
    """
    xml_tokens = _tokens(_xml_paragraph_texts(xml_bytes))
    extracted_tokens = _tokens(raw_texts)
    problems = []
    missing = xml_tokens - extracted_tokens
    phantom = extracted_tokens - xml_tokens
    for tok, cnt in list(missing.items())[:MAX_PROBLEM_SAMPLES]:
        problems.append(f"{label} 추출 누락 토큰 '{tok}' x{cnt} (텍스트박스·필드 의심)")
    if len(missing) > MAX_PROBLEM_SAMPLES:
        problems.append(f"... 외 {label} 누락 {len(missing) - MAX_PROBLEM_SAMPLES}종")
    for tok, cnt in list(phantom.items())[:MAX_PROBLEM_SAMPLES]:
        problems.append(f"{label} 원본에 없는 토큰 '{tok}' x{cnt}")
    return sum(xml_tokens.values()), problems


def verify_recall(filename, raw_texts):
    """본문(document.xml) recall 검증. 머리글·바닥글 파트는 별도로 대조한다."""
    with zipfile.ZipFile(filename) as zf:
        xml_bytes = zf.read("word/document.xml")
    return _verify_tokens(xml_bytes, raw_texts, "본문")


def collect_header_footer_parts(doc):
    """섹션별 머리글·바닥글 파트를 (라벨, 컨테이너) 목록으로.

    설정 가드: 첫 페이지·짝수 페이지 변형은 해당 설정이 켜진 경우에만 실제
    렌더되므로 그때만 포함한다. 이전 섹션에 연결된(linked) 파트는 정의 소유
    파트에서 이미 다루므로 건너뛰고, 같은 파트의 중복 등록은 partname으로 막는다.
    """
    parts = []
    seen = set()
    odd_even = bool(doc.settings.odd_and_even_pages_header_footer)
    for section in doc.sections:
        candidates = [("머리글", section.header), ("바닥글", section.footer)]
        if section.different_first_page_header_footer:
            candidates.append(("첫 페이지 머리글", section.first_page_header))
            candidates.append(("첫 페이지 바닥글", section.first_page_footer))
        if odd_even:
            candidates.append(("짝수 페이지 머리글", section.even_page_header))
            candidates.append(("짝수 페이지 바닥글", section.even_page_footer))
        for label, hf in candidates:
            if hf.is_linked_to_previous:
                continue
            partname = str(hf.part.partname)
            if partname in seen:
                continue
            seen.add(partname)
            parts.append((label, hf))
    return parts


def process_file(filename):
    import docx

    print(f"입력 파일: {filename}")
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".doc":
        print("오류: DOC(구형 바이너리)는 직접 처리할 수 없습니다.")
        print("  Word/LibreOffice에서 DOCX로 저장 후 입력하세요.")
        return
    if ext not in SUPPORTED_EXTENSIONS:
        print(f"오류: 지원하지 않는 파일 형식입니다: {ext}")
        print(f"지원 형식: {', '.join(SUPPORTED_EXTENSIONS)}")
        return
    if not os.path.exists(filename):
        print(f"오류: 파일을 찾을 수 없습니다: {filename}")
        return

    print("추출 중... (python-docx 본문 순서 보존 + XML 전수 recall 대조)")
    doc = docx.Document(filename)
    blocks, raw_texts, merged_cells, nested_tables = extract_blocks(doc)

    # 머리글·바닥글은 본문과 별개 XML 파트라 document.xml recall이 못 덮는다.
    # 파트별로 같은 추출(표 포함)·검증을 적용한다(첫/짝수 페이지 변형 포함).
    hf_sections = []  # (라벨, 블록 목록)
    hf_problems = []
    hf_token_total = 0
    sym_count = 0
    for label, hf in collect_header_footer_parts(doc):
        hf_blocks, hf_raws, hf_merged, hf_nested = extract_blocks(hf)
        merged_cells += hf_merged
        nested_tables += hf_nested
        count, part_problems = _verify_tokens(hf.part.blob, hf_raws, label)
        hf_token_total += count
        hf_problems.extend(part_problems)
        hf_xml = hf.part.blob.decode("utf-8", "replace")
        sym_count += hf_xml.count("<w:sym ") + hf_xml.count("<w:sym/")
        if hf_blocks:
            hf_sections.append((label, hf_blocks))

    with zipfile.ZipFile(filename) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8", "replace")
        footnote_text = _part_has_text(zf, "word/footnotes.xml")
        endnote_text = _part_has_text(zf, "word/endnotes.xml")
        has_comments = _part_has_text(zf, "word/comments.xml")
    images = document_xml.count("<w:drawing>") + document_xml.count("<w:drawing ")
    # w:sym(심볼 폰트 글리프)은 python-docx와 raw 집계 양쪽 다 못 읽어 recall
    # 대조로도 안 잡히는 대칭 침묵 유실이므로, 존재 자체를 감지해 거부한다.
    sym_count += document_xml.count("<w:sym ") + document_xml.count("<w:sym/")

    print(f"본문 블록 {len(blocks)}개 · 병합 셀 {merged_cells}곳")

    xml_token_count, problems = verify_recall(filename, raw_texts)
    print(
        f"전수 recall 검증(python-docx / XML 독립 집계): 본문 토큰 {xml_token_count}건"
        f" + 머리글·바닥글 토큰 {hf_token_total}건 대조"
    )

    blocking = list(problems) + hf_problems
    if nested_tables:
        blocking.append(f"중첩 표 {nested_tables}개: 마크다운으로 구조 보존 불가")
    if footnote_text:
        blocking.append("각주에 텍스트 존재: python-docx 미지원 영역(누락 위험)")
    if endnote_text:
        blocking.append("미주에 텍스트 존재: python-docx 미지원 영역(누락 위험)")
    if sym_count:
        blocking.append(f"심볼 글리프(w:sym) {sym_count}개: 양쪽 관점 모두 못 읽어 조용한 유실 위험")

    if blocking:
        print(f"경고: 검증·구조 문제 {len(blocking)}건")
        for line in blocking[:MAX_PROBLEM_SAMPLES + 3]:
            print(f"  - {line}")
        print("→ 검증 실패: 출력 파일을 만들지 않았습니다. 기존 티어(Upstage·LlamaParse)로 승격하세요.")
        return

    if merged_cells:
        print(f"[알림] 병합 셀 {merged_cells}곳: 값은 앵커 위치에 1회만 기록했습니다.")
    if has_comments:
        print("[알림] 문서에 검토 메모(comment)가 있습니다. 본문이 아니므로 출력에 포함하지 않았습니다.")
    if images:
        print(f"경고: 이미지 {images}개 존재. 그 안의 텍스트는 추출 범위 밖입니다.")
        print("  이미지 내 텍스트가 중요하면 Upstage(upstage_parse.py)로 교차검증하세요.")

    parts = list(blocks)
    if merged_cells:
        parts.append(f"병합 셀 {merged_cells}곳(값은 앵커 위치에 1회 기록)")
    for label, hf_blocks in hf_sections:
        for block in hf_blocks:
            if "\n" in block:  # 표 등 다행 블록은 라벨 줄을 앞세운다
                parts.append(f"{label}:")
                parts.append(block)
            else:
                parts.append(f"{label}: {block}")

    output_file = get_output_filename(filename)
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("\n\n".join(parts) + "\n")
    print(f"출력 파일: {output_file}")
    print("=== PASS: document.xml 전수 recall 일치 ===")


def main():
    try:
        import docx  # noqa: F401
    except ImportError as e:
        print("오류: python-docx 패키지를 찾을 수 없습니다.")
        print("설치 명령: pip install python-docx")
        print(f"상세: {e}")
        return

    positional = [a for a in sys.argv[1:] if not a.startswith("--")]
    if positional:
        process_file(positional[0])
        return

    supported_files = sorted(
        f for f in os.listdir(".")
        if os.path.isfile(f) and os.path.splitext(f)[1].lower() in SUPPORTED_EXTENSIONS
    )
    if not supported_files:
        print("오류: 현재 디렉토리에 지원되는 DOCX 파일이 없습니다.")
        print(f"지원 형식: {', '.join(SUPPORTED_EXTENSIONS)}")
        return
    if len(supported_files) == 1:
        process_file(supported_files[0])
        return
    print(f"DOCX 파일 {len(supported_files)}개 발견:")
    for i, f in enumerate(supported_files, 1):
        print(f"  {i}. {f}")
    print()
    for f in supported_files:
        process_file(f)
        print()


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"\n오류 발생: {e}")
        print("\n상세 정보:")
        traceback.print_exc()

    # 비-TTY(AI 에이전트·파이프·백그라운드)에서는 stdin이 EOF로 닫히지 않아
    # input()이 무한 블록되므로 isatty()로 가드 (다른 *_parse.py와 동일 패턴).
    if sys.stdin is not None and sys.stdin.isatty():
        try:
            input("\nEnter를 눌러 종료...")
        except EOFError:
            pass
